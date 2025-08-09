"""
Regulatory Compliance AI Assistant
-----------------------------------

This Streamlit application provides a user interface for ingesting
regulatory PDF documents, building a hybrid knowledge graph and
vector store, and performing retrieval‑augmented generation (RAG)
queries.  It improves upon earlier prototypes with a sleeker
tab‑based layout, support for uploading multiple documents at once,
precomputed executive summaries, clickable citations, and safer
prompt sizes to avoid hitting token or rate limits.  Chunks are
limited in size and retrieval context is truncated to maintain fast
response times.  Citations are rendered as expander elements so that
users can easily inspect the source text used to build an answer.

Environment variables required:

  OPENAI_API_KEY    – API key for the OpenAI chat API.
  NEO4J_URI         – Bolt(s) connection string for Neo4j (e.g. bolt+s://host:7687).
  NEO4J_USER        – Username for Neo4j.
  NEO4J_PASS        – Password for Neo4j.
  DATA_DIR          – Directory for SQLite and FAISS persistence (default: /content/data).

You can run this script with `streamlit run app_persisted.py`.
"""

from __future__ import annotations

import json
import os
import re
import sqlite3
import io
from datetime import datetime
from typing import Dict, List, Tuple, Optional

import streamlit as st
from py2neo import Graph, Node, Relationship  # type: ignore
from sentence_transformers import SentenceTransformer  # type: ignore
import faiss  # type: ignore
import numpy as np  # type: ignore
import fitz  # type: ignore  # PyMuPDF
from openai import OpenAI, RateLimitError  # type: ignore
from docx import Document  # type: ignore
from docx.shared import Pt  # type: ignore
from reportlab.lib.pagesizes import A4  # type: ignore
from reportlab.pdfgen import canvas  # type: ignore
from textwrap import wrap

# ---------------------------------------------------------------------------
# Configuration
# ---------------------------------------------------------------------------

# Hard limits to keep prompts within model token limits.  Chunks are
# intentionally small and context is truncated when generating answers.
MAX_CHARS_PER_CHUNK = 1800
MAX_SNIPPET_CHARS = 1200
MAX_TOTAL_CONTEXT_CHARS = 4800
TOP_K_RETRIEVE = 4

# Read environment variables with sane defaults.  Neo4j URI may be
# specified as neo4j+s:// (Aura) or bolt+s:// (py2neo prefers bolt).  We
# normalise this later.
DATA_DIR = os.getenv("DATA_DIR", "/content/data")
NEO4J_URI = os.getenv("NEO4J_URI", "bolt://localhost:7687")
NEO4J_USER = os.getenv("NEO4J_USER", "neo4j")
NEO4J_PASS = os.getenv("NEO4J_PASS", "password")

# Ensure persistence directory exists
os.makedirs(DATA_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# OpenAI client
# ---------------------------------------------------------------------------

# Use OpenAI SDK v1.  Credentials are read from environment.  See
# https://platform.openai.com/docs/api-reference for details.
client = OpenAI()

# Helper to safely call the chat completions API and handle rate or token
# errors gracefully.  Returns the message content or a warning string on
# failure.
def safe_chat(messages: List[Dict[str, str]], *, temperature: float = 0.0, model: str = "gpt-4o-mini") -> str:
    try:
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
        )
        return response.choices[0].message.content
    except RateLimitError as e:
        return f"⚠️ Rate limit / token error: {e}"
    except Exception as e:
        return f"⚠️ OpenAI error: {e}"

# Convenience wrapper to parse JSON from a language model response.  If
# parsing fails, returns an empty structure.  Handles fenced code blocks.
def llm_to_json(prompt: str) -> Dict[str, object]:
    raw = safe_chat([{"role": "user", "content": prompt}])
    try:
        return json.loads(raw)
    except Exception:
        try:
            clean = raw.strip("` ")
            return json.loads(clean)
        except Exception:
            return {"entities": [], "relationships": []}

# ---------------------------------------------------------------------------
# Neo4j integration
# ---------------------------------------------------------------------------

def normalise_bolt_uri(uri: str) -> str:
    """Convert neo4j or neo4j+s URIs to bolt or bolt+s for py2neo."""
    if uri.startswith("neo4j+s://"):
        return uri.replace("neo4j+s://", "bolt+s://")
    if uri.startswith("neo4j://"):
        return uri.replace("neo4j://", "bolt://")
    return uri

# Attempt to connect to Neo4j.  If connection fails, graph will be None.
_graph: Optional[Graph]
try:
    _graph = Graph(normalise_bolt_uri(NEO4J_URI), auth=(NEO4J_USER, NEO4J_PASS))
    # simple query to verify connectivity
    _ = _graph.run("RETURN 1").evaluate()
    GRAPH_CONNECTED = True
except Exception as e:
    _graph = None
    GRAPH_CONNECTED = False

# ---------------------------------------------------------------------------
# Persistence: SQLite + FAISS
# ---------------------------------------------------------------------------

# Create or open the vector persistence database
sql_path = os.path.join(DATA_DIR, "vectors.sqlite")
conn = sqlite3.connect(sql_path, check_same_thread=False)
conn.execute(
    """
    CREATE TABLE IF NOT EXISTS vectors (
        id INTEGER PRIMARY KEY,
        source TEXT,
        clause_id TEXT,
        text TEXT
    )
    """
)
conn.commit()

# Load or initialise FAISS index
faiss_path = os.path.join(DATA_DIR, "faiss.index")
dim = 384  # dimension for all-MiniLM-L6-v2 embeddings
try:
    faiss_index = faiss.read_index(faiss_path)
except Exception:
    faiss_index = faiss.IndexFlatL2(dim)

# Load cached vectors and rebuild FAISS if necessary
vector_store: List[Dict[str, object]] = []
cursor = conn.cursor()
cursor.execute("SELECT id, source, clause_id, text FROM vectors ORDER BY id")
rows = cursor.fetchall()
if rows:
    vector_store = [
        {
            "text": r[3],
            "metadata": {"source": r[1], "clause_id": r[2]},
        }
        for r in rows
    ]
    # rebuild FAISS if counts mismatch
    if faiss_index.ntotal != len(rows):
        faiss_index = faiss.IndexFlatL2(dim)
        # embed in batches
        embed_model = SentenceTransformer("all-MiniLM-L6-v2")
        for i in range(0, len(vector_store), 128):
            texts = [v["text"] for v in vector_store[i:i+128]]
            embs = embed_model.encode(texts)
            faiss_index.add(np.array(embs, dtype="float32"))
        faiss.write_index(faiss_index, faiss_path)

# ---------------------------------------------------------------------------
# Embeddings model
# ---------------------------------------------------------------------------

embed_model = SentenceTransformer("all-MiniLM-L6-v2")

# ---------------------------------------------------------------------------
# Utility functions
# ---------------------------------------------------------------------------

def pdf_to_chunks(pdf_bytes: bytes, regulation_name: str) -> List[Tuple[str, str]]:
    """Extract text from PDF and split into small chunks for embedding.

    PDFs are first split by 'Article <number>' headings when present to
    preserve semantic structure.  Each portion is then further split
    into paragraphs or character windows limited to MAX_CHARS_PER_CHUNK.
    """
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    full_text = "\n".join(page.get_text("text") for page in doc)
    # Try splitting by Article headings
    parts = re.split(r"(?=\bArticle\s+[0-9A-Za-z]+)", full_text)
    if len(parts) <= 1:
        parts = [full_text]
    chunks: List[Tuple[str, str]] = []
    for part in parts:
        cleaned = part.strip()
        if len(cleaned) < 100:
            continue
        # further split by blank lines
        paras = [p.strip() for p in re.split(r"\n\s*\n", cleaned) if p.strip()]
        buffer = ""
        for para in paras:
            if len(buffer) + len(para) + 2 <= MAX_CHARS_PER_CHUNK:
                buffer = (buffer + "\n\n" + para).strip()
            else:
                if buffer:
                    chunks.append((buffer, regulation_name))
                # break paragraphs longer than limit
                for i in range(0, len(para), MAX_CHARS_PER_CHUNK):
                    segment = para[i:i+MAX_CHARS_PER_CHUNK]
                    chunks.append((segment.strip(), regulation_name))
                buffer = ""
        if buffer:
            chunks.append((buffer, regulation_name))
    return chunks


def extract_entities_relationships(text: str, regulation_name: str) -> Dict[str, object]:
    """Use the LLM to extract entities and relationships from a text chunk."""
    prompt = f"""You are extracting knowledge graph data from regulatory text.
For the given text, output JSON with:
  entities: name, type (Regulation, Part, Clause, Term, Metric, Date)
  relationships: source, target, type (CONTAINS, DEFINES, REQUIRES, SUPERSEDES, REFERENCES, AFFECTS)
Return only valid JSON. Text to process:
```{text[:1500]}```"""
    data = llm_to_json(prompt)
    # Always ensure regulation node exists
    entities = data.get("entities", [])
    if not any(e.get("name") == regulation_name for e in entities):
        entities.append({"name": regulation_name, "type": "Regulation"})
        data["entities"] = entities
    return data


def load_into_neo4j(data: Dict[str, object]) -> None:
    """Store entities and relationships in Neo4j, if connected."""
    if not GRAPH_CONNECTED or _graph is None:
        return
    try:
        nodes: Dict[str, Node] = {}
        for ent in data.get("entities", []):
            label = ent.get("type", "Entity")
            name = ent.get("name")
            node = Node(label, name=name)
            _graph.merge(node, label, "name")
            nodes[name] = node
        for rel in data.get("relationships", []):
            src = rel.get("source")
            tgt = rel.get("target")
            rtype = rel.get("type", "RELATED")
            if src in nodes and tgt in nodes:
                _graph.merge(Relationship(nodes[src], rtype, nodes[tgt]))
    except Exception:
        pass


def detect_clause_id(text: str) -> Optional[str]:
    """Extract an Article/Clause identifier from a text chunk for citation."""
    m = re.search(r"(Article\s+[0-9A-Za-z]+)", text)
    return m.group(1) if m else None


def add_to_vector_store(text: str, metadata: Dict[str, str]) -> None:
    """Persist an embedded chunk in SQLite and FAISS."""
    # Save to SQLite
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO vectors(source, clause_id, text) VALUES (?, ?, ?)",
        (metadata.get("source"), metadata.get("clause_id"), text),
    )
    conn.commit()
    # Append to local store
    vector_store.append({"text": text, "metadata": metadata})
    # Add to FAISS index
    emb = embed_model.encode([text])
    faiss_index.add(np.array(emb, dtype="float32"))
    # Persist index to disk
    try:
        faiss.write_index(faiss_index, faiss_path)
    except Exception:
        pass


def semantic_search(query: str, top_k: int = TOP_K_RETRIEVE) -> List[Dict[str, object]]:
    """Retrieve top_k most similar chunks for a query over all sources."""
    if faiss_index.ntotal == 0:
        return []
    q_emb = embed_model.encode([query])
    distances, ids = faiss_index.search(np.array(q_emb, dtype="float32"), top_k)
    results = []
    for idx in ids[0]:
        if idx < len(vector_store):
            results.append(vector_store[idx])
    return results


def semantic_search_filtered(query: str, source: str, top_k: int = TOP_K_RETRIEVE) -> List[Dict[str, object]]:
    """Retrieve top_k chunks limited to a particular regulation."""
    if faiss_index.ntotal == 0:
        return []
    q_emb = embed_model.encode([query])
    distances, ids = faiss_index.search(np.array(q_emb, dtype="float32"), top_k * 5)
    results: List[Dict[str, object]] = []
    for idx in ids[0]:
        if idx < len(vector_store):
            meta = vector_store[idx]["metadata"]
            if meta.get("source") == source:
                results.append(vector_store[idx])
            if len(results) >= top_k:
                break
    return results


def build_context(snippets: List[Dict[str, object]]) -> str:
    """Assemble context string from snippets, truncating overall length."""
    ctx = ""
    for r in snippets:
        source = r["metadata"].get("source")
        cid = r["metadata"].get("clause_id") or "best match"
        snippet = r["text"].strip()
        snippet = snippet[:MAX_SNIPPET_CHARS] + (" …" if len(snippet) > MAX_SNIPPET_CHARS else "")
        part = f"Source ({source}, {cid}):\n{snippet}\n\n"
        if len(ctx) + len(part) > MAX_TOTAL_CONTEXT_CHARS:
            break
        ctx += part
    return ctx


def generate_answer(query: str, snippets: List[Dict[str, object]]) -> str:
    """Call the LLM to answer a query given retrieved snippets."""
    context = build_context(snippets)
    prompt = f"""You are a regulatory compliance assistant.

Question: {query}

Context:
{context}

Provide a concise answer. If comparing regulations, highlight differences.
Always cite sources in parentheses at the end of each sentence in the form (Regulation, Clause ID).
"""
    return safe_chat([{"role": "user", "content": prompt}])


def executive_summary_from_chunks(reg_name: str, chunks: List[str]) -> str:
    """Generate a short executive summary from sample chunks."""
    excerpts = "\n\n".join(chunk[:800] for chunk in chunks[:8])
    prompt = f"""You are a senior regulatory analyst. Draft a crisp executive summary for the regulation "{reg_name}".
Base your summary ONLY on the excerpts below. Cover scope, key obligations, reporting/governance and any effective dates.
Keep it under 160 words and avoid speculation.

EXCERPTS:
{excerpts}
"""
    return safe_chat([{"role": "user", "content": prompt}], temperature=0.2)


def basic_technical_overview(reg_name: str) -> Dict[str, int]:
    """Compute basic counts of graph and vector entities for a given regulation."""
    try:
        node_count = _graph.run("MATCH (n) RETURN count(n) AS c").evaluate() if GRAPH_CONNECTED else 0
        rel_count = _graph.run("MATCH ()-[r]-() RETURN count(r) AS c").evaluate() if GRAPH_CONNECTED else 0
        term_count = _graph.run("MATCH (t:Term) RETURN count(t) AS c").evaluate() if GRAPH_CONNECTED else 0
        clause_count = _graph.run("MATCH (c:Clause) RETURN count(c) AS c").evaluate() if GRAPH_CONNECTED else 0
    except Exception:
        node_count = rel_count = term_count = clause_count = 0
    indexed = sum(1 for v in vector_store if v["metadata"].get("source") == reg_name)
    return {
        "nodes": int(node_count or 0),
        "relationships": int(rel_count or 0),
        "terms": int(term_count),
        "clauses": int(clause_count),
        "indexed_clauses": indexed,
    }


# ---------------------------------------------------------------------------
# Streamlit Interface
# ---------------------------------------------------------------------------

st.sidebar.title("Regulatory AI Assistant")
st.sidebar.markdown(
    "This tool ingests regulatory PDFs, builds a knowledge graph and vector store, and lets you compare rules and query them."
)

st.title("Regulatory Compliance AI Assistant")

# Use tabs for a sleek UI
tabs = st.tabs([
    "Ingest", "Summaries", "Compare", "Batch Report", "Q&A", "Settings"
])

exec_summary_cache: Dict[str, str] = st.session_state.setdefault("exec_summary_cache", {})

# ---------------------------------------------------------------------------
# Tab: Ingest
# ---------------------------------------------------------------------------
with tabs[0]:
    st.header("📥 Ingest Regulations")
    st.write(
        "Upload one or more PDF files and assign a regulation name. This will extract text, build embeddings, and optionally load entities into Neo4j."
    )
    uploaded_files = st.file_uploader(
        "Upload regulation PDFs", type="pdf", accept_multiple_files=True
    )
    reg_name = st.text_input(
        "Regulation Name (e.g., CRR III, Internal Policy)", key="reg_name"
    )
    if uploaded_files and reg_name and st.button("Process & Ingest"):
        with st.spinner("Processing uploaded documents..."):
            sample_chunks: List[str] = []
            for file in uploaded_files:
                chunks = pdf_to_chunks(file.read(), reg_name)
                for text, _src in chunks:
                    # extract KG and store
                    data = extract_entities_relationships(text, reg_name)
                    load_into_neo4j(data)
                    cid = detect_clause_id(text)
                    add_to_vector_store(text, {"source": reg_name, "clause_id": cid})
                    sample_chunks.append(text)
            # Precompute and cache executive summary
            exec_summary_cache[reg_name] = executive_summary_from_chunks(reg_name, sample_chunks)
        st.success(f"Successfully ingested {len(uploaded_files)} file(s) for {reg_name}.")


# ---------------------------------------------------------------------------
# Tab: Summaries
# ---------------------------------------------------------------------------
with tabs[1]:
    st.header("📄 Summaries & Technical Overview")
    regs = sorted({v["metadata"].get("source") for v in vector_store if v["metadata"].get("source")})
    if not regs:
        st.info("No regulations ingested yet. Please upload a PDF under the Ingest tab.")
    else:
        col1, col2 = st.columns(2)
        with col1:
            selected_reg = st.selectbox("Select a regulation", regs, key="summary_reg")
            if st.button("Generate / Refresh Executive Summary"):
                with st.spinner("Generating executive summary..."):
                    sample = [v["text"] for v in vector_store if v["metadata"].get("source") == selected_reg][:25]
                    exec_summary_cache[selected_reg] = executive_summary_from_chunks(selected_reg, sample)
            summary = exec_summary_cache.get(selected_reg)
            if summary:
                st.subheader(f"Executive Summary — {selected_reg}")
                st.write(summary)
            else:
                st.info("No executive summary available. Please refresh.")
        with col2:
            stat_reg = st.selectbox("Regulation for technical details", regs, key="tech_reg")
            if st.button("Show Technical Overview"):
                with st.spinner("Gathering stats..."):
                    stats = basic_technical_overview(stat_reg)
                st.subheader(f"Technical Overview — {stat_reg}")
                st.write(
                    f"Nodes: {stats['nodes']} | Relationships: {stats['relationships']} | "
                    f"Terms: {stats['terms']} | Clauses (graph): {stats['clauses']} | "
                    f"Indexed Clauses (vectors): {stats['indexed_clauses']}"
                )


# ---------------------------------------------------------------------------
# Tab: Compare
# ---------------------------------------------------------------------------
with tabs[2]:
    st.header("⚖️ Compare Regulations")
    regs = sorted({v["metadata"].get("source") for v in vector_store if v["metadata"].get("source")})
    if len(regs) < 2:
        st.info("At least two regulations must be ingested to compare.")
    else:
        col1, col2 = st.columns(2)
        with col1:
            reg_a = st.selectbox("Regulation A", regs, key="compare_reg_a")
        with col2:
            reg_b = st.selectbox("Regulation B", regs, index=1 if len(regs) > 1 else 0, key="compare_reg_b")
        topic = st.text_input(
            "Topic or term to compare (e.g., 'Liquidity Coverage Ratio')",
            key="compare_topic",
        )
        if st.button("Compare") and topic:
            with st.spinner("Retrieving relevant clauses..."):
                hits_a = semantic_search_filtered(topic, reg_a, top_k=TOP_K_RETRIEVE)
                hits_b = semantic_search_filtered(topic, reg_b, top_k=TOP_K_RETRIEVE)
                hit_a = hits_a[0] if hits_a else None
                hit_b = hits_b[0] if hits_b else None
            if not hit_a or not hit_b:
                st.warning("No matching clauses found for one or both regulations.")
            else:
                st.subheader("Matched Clauses")
                c1, c2 = st.columns(2)
                with c1:
                    cid_a = hit_a["metadata"].get("clause_id", "best match")
                    st.caption(f"{reg_a} — {cid_a}")
                    st.write(hit_a["text"])
                with c2:
                    cid_b = hit_b["metadata"].get("clause_id", "best match")
                    st.caption(f"{reg_b} — {cid_b}")
                    st.write(hit_b["text"])
                st.subheader("Differences (approximate)")
                # Render diff as simple side by side difference: highlight differences inline
                diff_text = html_diff(reg_a, hit_a["text"], reg_b, hit_b["text"])
                st.components.v1.html(diff_text, height=420, scrolling=True)
                st.subheader("Summary with Citations")
                summary = generate_answer(
                    f"Compare how these two regulations address: {topic}", [hit_a, hit_b]
                )
                st.write(summary)


# ---------------------------------------------------------------------------
# Tab: Batch Report
# ---------------------------------------------------------------------------
with tabs[3]:
    st.header("🗂️ Batch Comparisons & Report")
    regs = sorted({v["metadata"].get("source") for v in vector_store if v["metadata"].get("source")})
    if len(regs) < 2:
        st.info("At least two regulations must be ingested to build a report.")
    else:
        col1, col2 = st.columns(2)
        with col1:
            batch_a = st.selectbox("Regulation A (batch)", regs, key="batch_reg_a")
        with col2:
            batch_b = st.selectbox("Regulation B (batch)", regs, index=1 if len(regs) > 1 else 0, key="batch_reg_b")
        topics_input = st.text_area(
            "Topics (one per line)",
            "Liquidity Coverage Ratio\nOperational Risk\nReporting Frequency",
            key="batch_topics",
        )
        include_exec = st.checkbox("Include Executive Summaries", True)
        if st.button("Generate Report") and topics_input.strip():
            topics_list = [t.strip() for t in topics_input.splitlines() if t.strip()]
            results: List[Dict[str, object]] = []
            exec_a: Optional[str] = None
            exec_b: Optional[str] = None
            with st.spinner("Compiling comparisons..."):
                # Precompute exec summaries if requested
                if include_exec:
                    exec_a = exec_summary_cache.get(batch_a)
                    if not exec_a:
                        samples_a = [v["text"] for v in vector_store if v["metadata"].get("source") == batch_a][:25]
                        exec_a = executive_summary_from_chunks(batch_a, samples_a)
                        exec_summary_cache[batch_a] = exec_a
                    exec_b = exec_summary_cache.get(batch_b)
                    if not exec_b:
                        samples_b = [v["text"] for v in vector_store if v["metadata"].get("source") == batch_b][:25]
                        exec_b = executive_summary_from_chunks(batch_b, samples_b)
                        exec_summary_cache[batch_b] = exec_b
                for topic_name in topics_list:
                    h_a = semantic_search_filtered(topic_name, batch_a, top_k=TOP_K_RETRIEVE)
                    h_b = semantic_search_filtered(topic_name, batch_b, top_k=TOP_K_RETRIEVE)
                    if h_a and h_b:
                        hit_a = h_a[0]
                        hit_b = h_b[0]
                        results.append({
                            "topic": topic_name,
                            "a": hit_a,
                            "b": hit_b,
                            "summary": generate_answer(
                                f"Compare how these two regulations treat: {topic_name}", [hit_a, hit_b]
                            ),
                        })
            if results:
                st.success(f"Generated report for {len(results)} topics.")
                # Provide downloads
                def build_docx() -> bytes:
                    doc = Document()
                    doc.add_heading("Multi‑topic Regulatory Comparison Report", 0)
                    doc.add_paragraph(
                        f"Generated: {datetime.utcnow().isoformat(timespec='seconds')}Z"
                    )
                    doc.add_paragraph(f"Regulations compared: {batch_a} vs {batch_b}")
                    if include_exec:
                        doc.add_heading("Executive Summaries", level=1)
                        doc.add_heading(batch_a, level=2)
                        doc.add_paragraph(exec_a or "(none)")
                        doc.add_heading(batch_b, level=2)
                        doc.add_paragraph(exec_b or "(none)")
                    doc.add_heading("Comparisons", level=1)
                    for item in results:
                        doc.add_heading(item["topic"], level=2)
                        doc.add_paragraph(item["summary"])
                        # Clause text and citations
                        a_meta = item["a"]["metadata"]
                        b_meta = item["b"]["metadata"]
                        doc.add_heading(f"{batch_a} — {a_meta.get('clause_id', 'best match')}", level=3)
                        doc.add_paragraph(item["a"]["text"])
                        doc.add_heading(f"{batch_b} — {b_meta.get('clause_id', 'best match')}", level=3)
                        doc.add_paragraph(item["b"]["text"])
                        doc.add_heading("Citations", level=3)
                        doc.add_paragraph(f"- ({batch_a}, {a_meta.get('clause_id', 'best match')})")
                        doc.add_paragraph(f"- ({batch_b}, {b_meta.get('clause_id', 'best match')})")
                    bio = io.BytesIO()
                    doc.save(bio)
                    bio.seek(0)
                    return bio.getvalue()

                def build_pdf() -> bytes:
                    buffer = io.BytesIO()
                    c = canvas.Canvas(buffer, pagesize=A4)
                    W, H = A4
                    margin = 40
                    line_h = 14
                    def draw_lines(title: str, body: str, y: float, bold: bool = False) -> float:
                        c.setFont("Helvetica-Bold" if bold else "Helvetica", 12 if bold else 10)
                        text = title if bold else body
                        for line in wrap(text, 110):
                            if y < 60:
                                c.showPage()
                                y = H - 60
                                c.setFont("Helvetica", 10)
                            c.drawString(margin, y, line)
                            y -= line_h
                        return y - (6 if bold else 2)
                    c.setFont("Helvetica-Bold", 16)
                    c.drawString(margin, H - margin, "Multi‑topic Regulatory Comparison Report")
                    c.setFont("Helvetica", 9)
                    c.drawString(margin, H - margin - 18, f"Generated: {datetime.utcnow().isoformat(timespec='seconds')}Z")
                    c.drawString(margin, H - margin - 32, f"Regulations compared: {batch_a} vs {batch_b}")
                    y = H - margin - 54
                    if include_exec:
                        y = draw_lines("Executive Summaries", "", y, bold=True)
                        y = draw_lines(batch_a, exec_a or "(none)", y)
                        y = draw_lines(batch_b, exec_b or "(none)", y)
                    y = draw_lines("Comparisons", "", y, bold=True)
                    for item in results:
                        y = draw_lines(item["topic"], "", y, bold=True)
                        y = draw_lines("Summary:", item["summary"], y)
                        a_meta = item["a"]["metadata"]
                        b_meta = item["b"]["metadata"]
                        y = draw_lines(f"{batch_a} — {a_meta.get('clause_id', 'best match')}", item["a"]["text"], y)
                        y = draw_lines(f"{batch_b} — {b_meta.get('clause_id', 'best match')}", item["b"]["text"], y)
                        y = draw_lines("Citations:", f"({batch_a}, {a_meta.get('clause_id', 'best match')}) | ({batch_b}, {b_meta.get('clause_id', 'best match')})", y)
                    c.showPage()
                    c.save()
                    buffer.seek(0)
                    return buffer.getvalue()
                colA, colB = st.columns(2)
                with colA:
                    st.download_button(
                        "Download DOCX",
                        data=build_docx(),
                        file_name=f"comparison_{batch_a}_vs_{batch_b}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                with colB:
                    st.download_button(
                        "Download PDF",
                        data=build_pdf(),
                        file_name=f"comparison_{batch_a}_vs_{batch_b}.pdf",
                        mime="application/pdf",
                    )
            else:
                st.warning("No matches found for the specified topics.")


# ---------------------------------------------------------------------------
# Tab: Q&A
# ---------------------------------------------------------------------------
with tabs[4]:
    st.header("💬 Ask a Question")
    regs = sorted({v["metadata"].get("source") for v in vector_store if v["metadata"].get("source")})
    question = st.text_input("Enter your compliance question", key="qa_question")
    if st.button("Get Answer") and question:
        with st.spinner("Searching and generating response..."):
            hits = semantic_search(question, top_k=TOP_K_RETRIEVE)
            answer = generate_answer(question, hits)
        st.subheader("Answer with Citations")
        st.write(answer)
        st.subheader("Source Clauses")
        if hits:
            for hit in hits:
                citation = hit["metadata"].get("clause_id", "best match")
                with st.expander(f"{hit['metadata'].get('source')} — {citation}"):
                    st.write(hit["text"])
        else:
            st.write("No matching clauses found.")


# ---------------------------------------------------------------------------
# Tab: Settings
# ---------------------------------------------------------------------------
with tabs[5]:
    st.header("⚙️ Settings & Info")
    st.write("Configure advanced options and view system status.")
    st.write(
        f"**Graph connected:** {'Yes' if GRAPH_CONNECTED else 'No'} | "
        f"**Indexed chunks:** {faiss_index.ntotal} | **Regulations:** {len({v['metadata'].get('source') for v in vector_store})}"
    )
    st.write(
        "**Neo4j URI:** ", normalise_bolt_uri(NEO4J_URI), "\n",
        "**Data directory:** ", DATA_DIR
    )

    if st.button("Clear persisted data"):
        # Danger: remove SQLite and FAISS index
        try:
            conn.close()
            os.remove(sql_path)
            os.remove(faiss_path)
            st.success("Persisted data cleared. Please restart the app.")
        except Exception as e:
            st.error(f"Failed to clear data: {e}")

    st.write("\nThis is a proof‑of‑concept. For production use, consider hosting on a cloud service (AWS, Azure, GCP) with persistent storage and adding user authentication.")
